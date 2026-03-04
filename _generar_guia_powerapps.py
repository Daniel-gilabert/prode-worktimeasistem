from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos de pagina ──
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def titulo_portada(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x00, 0x47, 0xAB)
    return p

def subtitulo_portada(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p

def seccion(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x47, 0xAB)
    return h

def parrafo(doc, texto, negrita=False, sangria=False):
    p = doc.add_paragraph()
    if sangria:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(texto)
    run.bold = negrita
    run.font.size = Pt(10.5)
    return p

def nota(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run("NOTA: " + texto)
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
    run.font.size = Pt(10)
    return p

def lista(doc, items, numerada=False):
    for item in items:
        style = 'List Number' if numerada else 'List Bullet'
        p = doc.add_paragraph(item, style=style)
        p.paragraph_format.left_indent = Inches(0.3)
        for run in p.runs:
            run.font.size = Pt(10.5)

def tabla(doc, headers, filas, col_widths=None, header_color='002F6C'):
    t = doc.add_table(rows=len(filas)+1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        set_cell_bg(cell, header_color)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
    for ri, fila in enumerate(filas):
        bg = 'F0F4FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(fila):
            cell = t.rows[ri+1].cells[ci]
            cell.text = val
            set_cell_bg(cell, bg)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph('')
    return t

def separador(doc):
    p = doc.add_paragraph('─' * 80)
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ════════════════════════════════════════
# PORTADA
# ════════════════════════════════════════
doc.add_paragraph('')
doc.add_paragraph('')
titulo_portada(doc, 'WorkTimeAsistem PRODE')
titulo_portada(doc, 'Guia de Migración a Microsoft Power Apps')
doc.add_paragraph('')
subtitulo_portada(doc, 'Guia tecnica completa — paso a paso')
subtitulo_portada(doc, f'Version 1.0 — {datetime.datetime.now().strftime("%d/%m/%Y")}')
subtitulo_portada(doc, 'Elaborado por Daniel Gilabert Cantero para Fundacion PRODE')
doc.add_paragraph('')
doc.add_paragraph('')

p_aviso = doc.add_paragraph()
p_aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_aviso.add_run(
    'Este documento describe como trasladar el sistema actual (Python/Streamlit + Supabase)\n'
    'a Microsoft Power Apps con Dataverse o SharePoint como backend,\n'
    'aprovechando la licencia Microsoft 365 existente en PRODE.'
)
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_page_break()

# ════════════════════════════════════════
# INDICE
# ════════════════════════════════════════
seccion(doc, 'Indice de contenidos')
indice = [
    '1.  Por que migrar a Power Apps',
    '2.  Comparativa: sistema actual vs Power Apps',
    '3.  Requisitos previos',
    '4.  Arquitectura objetivo en Power Apps',
    '5.  Fase 1 — Preparacion del entorno Microsoft 365',
    '6.  Fase 2 — Modelo de datos: Dataverse vs SharePoint',
    '7.  Fase 3 — Migracion de datos desde Supabase',
    '8.  Fase 4 — Construccion de la app en Power Apps',
    '9.  Fase 5 — Logica de negocio con Power Automate',
    '10. Fase 6 — Informes PDF y Excel con Power Automate',
    '11. Fase 7 — Integracion con Power BI',
    '12. Fase 8 — Seguridad, roles y permisos en Power Platform',
    '13. Fase 9 — Pruebas y validacion',
    '14. Fase 10 — Despliegue y publicacion',
    '15. Cronograma orientativo',
    '16. Costes estimados',
    '17. Riesgos y mitigaciones',
    '18. Glosario',
]
lista(doc, indice)
doc.add_page_break()

# ════════════════════════════════════════
# 1. POR QUE MIGRAR
# ════════════════════════════════════════
seccion(doc, '1. Por que migrar a Power Apps')
parrafo(doc,
    'La migracion a Power Apps tiene sentido si PRODE ya dispone de licencias Microsoft 365, '
    'porque permite eliminar costes externos (Streamlit Cloud, Supabase), mejorar la integracion '
    'con Teams, SharePoint y Outlook, y reducir la dependencia de codigo Python que requiere '
    'mantenimiento tecnico especializado.')
doc.add_paragraph('')
parrafo(doc, 'Ventajas principales:', negrita=True)
lista(doc, [
    'Sin servidor propio: Microsoft gestiona la infraestructura.',
    'Licencias ya pagadas: Power Apps esta incluido en Microsoft 365 Business.',
    'Integracion nativa con Teams, Outlook, SharePoint y Power BI.',
    'No-code / Low-code: cualquier persona con formacion basica puede mantener la app.',
    'Seguridad gestionada por Azure AD: los usuarios se autentican con sus cuentas corporativas @prode.es sin login adicional.',
    'Power Automate reemplaza scripts Python para logica de negocio y envio de correos.',
    'Versionado de la app integrado en Power Apps Studio.',
])
doc.add_paragraph('')
nota(doc, 'Si PRODE no tiene licencias Power Apps per user, el coste adicional es aprox. 20 EUR/usuario/mes. '
          'Verificar con el area IT si las licencias actuales incluyen Power Apps.')

doc.add_page_break()

# ════════════════════════════════════════
# 2. COMPARATIVA
# ════════════════════════════════════════
seccion(doc, '2. Comparativa: sistema actual vs Power Apps')
tabla(doc,
    ['Aspecto', 'Sistema actual (Streamlit + Supabase)', 'Power Apps + Dataverse'],
    [
        ['Backend', 'Supabase (PostgreSQL)', 'Dataverse o SharePoint Online'],
        ['Frontend', 'Python Streamlit', 'Power Apps Canvas App'],
        ['Logica', 'Python (calculo_service.py etc.)', 'Power Automate + formulas Power Apps'],
        ['Autenticacion', 'Login manual por email', 'Azure AD (SSO con cuenta @prode.es)'],
        ['Roles', 'Campo rol en tabla empleados', 'Grupos de seguridad Azure AD o roles Dataverse'],
        ['PDF/Excel', 'ReportLab + OpenPyXL', 'Power Automate con plantillas Word/Excel'],
        ['Despliegue', 'Streamlit Cloud + .env', 'Power Apps Service (nube Microsoft)'],
        ['Coste', 'Gratis (Supabase free) + dev time', 'Incluido en M365 o +20 EUR/user/mes'],
        ['Mantenimiento', 'Requiere Python dev', 'Low-code, mantenible por admin IT'],
        ['Integracion Power BI', 'Via PostgreSQL o API REST', 'Nativa (mismo ecosistema)'],
        ['Teams', 'No integrado', 'App embebible en Teams'],
    ],
    col_widths=[1.5, 2.5, 2.5]
)

doc.add_page_break()

# ════════════════════════════════════════
# 3. REQUISITOS PREVIOS
# ════════════════════════════════════════
seccion(doc, '3. Requisitos previos')
parrafo(doc, 'Antes de iniciar la migracion, verificar:', negrita=True)
lista(doc, [
    'Licencias Microsoft 365: al menos Business Basic o Business Standard para todos los usuarios de la app.',
    'Licencias Power Apps: incluidas en M365 o adquirir Power Apps per user (plan Premium para Dataverse).',
    'Acceso a Power Platform Admin Center: https://admin.powerplatform.microsoft.com/',
    'Tenant de Azure AD configurado con los correos @prode.es.',
    'Power BI Pro o Premium por usuario si se quiere publicar dashboards.',
    'Permisos de administrador global en M365 o al menos administrador de Power Platform.',
    'Export de datos de Supabase: volcado CSV o JSON de todas las tablas antes de iniciar.',
])
doc.add_paragraph('')

parrafo(doc, 'Personas clave del proyecto:', negrita=True)
tabla(doc,
    ['Rol en proyecto', 'Responsabilidad', 'Persona sugerida'],
    [
        ['Owner / Patron', 'Toma de decisiones y presupuesto', 'Direccion PRODE'],
        ['Administrador Power Platform', 'Configuracion entorno, licencias, seguridad', 'IT PRODE'],
        ['Desarrollador Power Apps', 'Construccion de pantallas y formulas', 'Daniel Gilabert / consultor externo'],
        ['Experto de negocio', 'Validar logica de calculos y flujos', 'Daniel Gilabert'],
        ['Usuarios piloto', 'Probar app antes del despliegue total', '2-3 responsables de departamento'],
    ],
    col_widths=[1.8, 2.5, 2.0]
)

doc.add_page_break()

# ════════════════════════════════════════
# 4. ARQUITECTURA OBJETIVO
# ════════════════════════════════════════
seccion(doc, '4. Arquitectura objetivo en Power Apps')
parrafo(doc,
    'La arquitectura de destino usa exclusivamente servicios Microsoft, eliminando Supabase y Streamlit:')
doc.add_paragraph('')

tabla(doc,
    ['Capa', 'Tecnologia', 'Equivalente actual'],
    [
        ['UI / Frontend', 'Power Apps Canvas App', 'Streamlit (app.py + ui/*.py)'],
        ['Logica de negocio', 'Power Automate (Cloud Flows)', 'services/*.py (Python)'],
        ['Base de datos', 'Microsoft Dataverse', 'Supabase (PostgreSQL)'],
        ['Almacenamiento ficheros', 'SharePoint Online / Azure Blob', 'Supabase Storage'],
        ['Autenticacion', 'Azure Active Directory', 'Login manual por email'],
        ['Informes PDF', 'Power Automate + Word Online', 'ReportLab (informe_pdf_service.py)'],
        ['Informes Excel', 'Power Automate + Excel Online', 'OpenPyXL (informe_excel_service.py)'],
        ['Analisis / BI', 'Power BI (conector Dataverse nativo)', 'Power BI + export Excel manual'],
        ['Notificaciones', 'Power Automate + Outlook', 'No implementado actualmente'],
        ['Auditoria', 'Dataverse audit log + tabla auditoria', 'Tabla auditoria en Supabase'],
    ],
    col_widths=[1.5, 2.0, 2.5]
)

doc.add_page_break()

# ════════════════════════════════════════
# 5. FASE 1 — ENTORNO M365
# ════════════════════════════════════════
seccion(doc, '5. Fase 1 — Preparacion del entorno Microsoft 365')
seccion(doc, '5.1 Crear entorno de Power Platform', nivel=2)
lista(doc, [
    'Ir a https://admin.powerplatform.microsoft.com/',
    'Entornos > + Nuevo',
    'Nombre: "WorkTimeAsistem-PRODE-PRD" (produccion) y "WorkTimeAsistem-PRODE-DEV" (desarrollo)',
    'Tipo: Produccion (para PRD) y Sandbox (para DEV)',
    'Region: Europa (para cumplir RGPD)',
    'Activar Dataverse: Si',
    'Moneda: EUR, Idioma: Espanol',
    'Crear entorno y esperar 5-10 min a que se provisione.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '5.2 Asignar licencias a usuarios', nivel=2)
lista(doc, [
    'Centro de administracion M365: https://admin.microsoft.com/',
    'Usuarios > Usuarios activos > seleccionar usuario',
    'Licencias y aplicaciones > activar Power Apps (si disponible en el plan)',
    'Repetir para todos los usuarios que accedan a la app.',
    'Los empleados sin acceso a la app NO necesitan licencia Power Apps.',
], numerada=True)

doc.add_page_break()

# ════════════════════════════════════════
# 6. FASE 2 — MODELO DE DATOS
# ════════════════════════════════════════
seccion(doc, '6. Fase 2 — Modelo de datos: Dataverse vs SharePoint')
parrafo(doc, 'Hay dos opciones para el backend. Se recomienda Dataverse si hay licencias Premium:', negrita=True)
doc.add_paragraph('')

tabla(doc,
    ['Criterio', 'Dataverse', 'SharePoint Online'],
    [
        ['Licencia', 'Power Apps Premium (~20 EUR/user/mes)', 'Incluida en M365'],
        ['Capacidad', '2 GB por defecto, escalable', 'Limitada por sitio SharePoint'],
        ['Relaciones', 'Relaciones reales (FK)', 'Solo lookups simples'],
        ['Seguridad filas', 'Row-level security nativo', 'Permisos por lista, mas complejo'],
        ['Rendimiento', 'Alto, optimizado para apps', 'Aceptable hasta ~5.000 filas'],
        ['Integracion Power BI', 'Conector Dataverse nativo, tiempo real', 'Conector SharePoint, mas lento'],
        ['Auditoria', 'Log de auditoria nativo', 'Manual o via Power Automate'],
        ['Recomendacion', 'RECOMENDADO para PRODE', 'Alternativa si no hay presupuesto'],
    ],
    col_widths=[1.8, 2.0, 2.3]
)
doc.add_paragraph('')

seccion(doc, '6.1 Tablas Dataverse equivalentes a Supabase', nivel=2)
tabla(doc,
    ['Tabla Supabase', 'Tabla Dataverse', 'Cambios'],
    [
        ['empleados', 'prode_Empleado', 'Igual. Campo "rol" como Choice (lista de opciones).'],
        ['festivos_locales', 'prode_FestivoLocal', 'Igual. Relacion con prode_Empleado (responsable).'],
        ['festivos_empleado', 'prode_FestivoEmpleado', 'Relacion M:M nativa en Dataverse.'],
        ['incidencias', 'prode_Incidencia', 'Campo "tipo" como Choice: VACACIONES/BAJA/PERMISO.'],
        ['historico_resumenes', 'prode_HistoricoResumen', 'Mismos campos. Relacion con prode_Empleado.'],
        ['panel_acceso', 'prode_PanelAcceso', 'Tabla de correos autorizados.'],
        ['auditoria', 'prode_Auditoria', 'Log de acciones. Campo usuario, accion, fecha, detalle.'],
    ],
    col_widths=[1.8, 1.8, 2.5]
)

seccion(doc, '6.2 Crear tablas en Dataverse', nivel=2)
lista(doc, [
    'Ir a https://make.powerapps.com/ > seleccionar entorno PRODE-DEV',
    'Dataverse > Tablas > + Nueva tabla',
    'Para cada tabla: definir nombre, nombre de visualizacion y columnas.',
    'Columnas de tipo Choice para campos como "rol" e "incidencia.tipo".',
    'Columnas de tipo Lookup para relaciones entre tablas (ej: empleado.responsable_id).',
    'Activar "Track changes" en todas las tablas para auditoria.',
    'Configurar permisos de tabla: a nivel entorno y a nivel fila (row-level security).',
], numerada=True)

doc.add_page_break()

# ════════════════════════════════════════
# 7. FASE 3 — MIGRACION DE DATOS
# ════════════════════════════════════════
seccion(doc, '7. Fase 3 — Migracion de datos desde Supabase a Dataverse')
parrafo(doc, 'La migracion de datos es un paso critico. Realizarla en entorno DEV primero.', negrita=True)
doc.add_paragraph('')

seccion(doc, '7.1 Exportar datos de Supabase', nivel=2)
lista(doc, [
    'Ir a https://supabase.com/dashboard/project/gqfiarxccbaznjxispsv',
    'Table editor > seleccionar tabla > "Export as CSV"',
    'Exportar todas las tablas: empleados, festivos_locales, festivos_empleado, incidencias, historico_resumenes, auditoria.',
    'Guardar los CSV en una carpeta local o SharePoint.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '7.2 Importar datos a Dataverse', nivel=2)
parrafo(doc, 'Opcion A — Importacion directa desde Power Apps (recomendada para datos iniciales):', negrita=True)
lista(doc, [
    'En make.powerapps.com > Dataverse > Tablas > seleccionar tabla > Importar > Importar datos de Excel.',
    'Subir el CSV exportado de Supabase (convertir a XLSX si es necesario).',
    'Mapear columnas del CSV a columnas de Dataverse.',
    'Verificar datos importados antes de continuar.',
], numerada=True)
doc.add_paragraph('')

parrafo(doc, 'Opcion B — Script Python de migracion (recomendada para datos relacionales):', negrita=True)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    'import pandas as pd\n'
    'from supabase import create_client\n'
    'from msal import ConfidentialClientApplication\n'
    'import requests\n\n'
    '# 1. Leer de Supabase\n'
    'sb = create_client(SUPABASE_URL, SUPABASE_KEY)\n'
    'empleados = sb.table("empleados").select("*").execute().data\n\n'
    '# 2. Autenticar en Dataverse\n'
    'app = ConfidentialClientApplication(CLIENT_ID, CLIENT_SECRET, AUTHORITY)\n'
    'token = app.acquire_token_for_client(["https://orgXXX.crm4.dynamics.com/.default"])\n\n'
    '# 3. Insertar en Dataverse via API\n'
    'headers = {"Authorization": f"Bearer {token[\'access_token\']}", "Content-Type": "application/json"}\n'
    'for emp in empleados:\n'
    '    requests.post(f"{DATAVERSE_URL}/api/data/v9.2/prode_empleados", json=emp, headers=headers)\n'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'
doc.add_paragraph('')

nota(doc, 'Instalar libreria MSAL: pip install msal. Obtener CLIENT_ID y CLIENT_SECRET registrando '
          'una app en Azure AD (portal.azure.com > Azure Active Directory > App registrations).')

doc.add_page_break()

# ════════════════════════════════════════
# 8. FASE 4 — CONSTRUCCION EN POWER APPS
# ════════════════════════════════════════
seccion(doc, '8. Fase 4 — Construccion de la app en Power Apps Canvas')
parrafo(doc,
    'Power Apps Canvas permite disenar la app "lienzo" con total libertad visual, similar a Streamlit. '
    'Es la opcion recomendada para replicar la UI actual con mayor flexibilidad que Model-driven apps.')
doc.add_paragraph('')

seccion(doc, '8.1 Crear la app Canvas', nivel=2)
lista(doc, [
    'Ir a https://make.powerapps.com/ > seleccionar entorno PRODE-DEV',
    'Apps > + Nueva app > Canvas',
    'Formato: Telefono (vertical) o Tableta (horizontal, recomendado para tabla de resumen).',
    'Nombrar: "WorkTimeAsistem PRODE".',
    'Conectar origenes de datos: + Agregar datos > Microsoft Dataverse > seleccionar todas las tablas prode_*.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '8.2 Pantallas a crear (equivalencia con UI actual)', nivel=2)
tabla(doc,
    ['Pantalla Power Apps', 'Equivalente Streamlit', 'Componentes principales'],
    [
        ['Pantalla Login', 'ui/login.py', 'Logo PRODE, boton "Iniciar sesion con Microsoft" (Azure AD)'],
        ['Pantalla Inicio', 'app.py (routing)', 'Menu segun rol: icono semaforo, resumen, config, admin'],
        ['Pantalla Resumen', 'ui/resumen.py', 'Galeria de empleados con columnas y semaforo de color'],
        ['Pantalla Configuracion', 'ui/configuracion.py', 'Tabs: Jornadas, Festivos, Incidencias con formularios'],
        ['Pantalla Panel Semaforo', 'ui/panel_responsables.py', 'Tarjetas por departamento con indicadores'],
        ['Pantalla Historico', 'seccion historico en resumen', 'Galeria de archivos + graficas (Power BI embed)'],
        ['Pantalla Panel Control', 'ui/panel_control.py', 'Lista empleados + formulario edicion rol/dept'],
        ['Pantalla Auditoria', 'tabla auditoria', 'Galeria con filtros fecha/usuario/accion'],
    ],
    col_widths=[2.0, 1.8, 2.5]
)
doc.add_paragraph('')

seccion(doc, '8.3 Autenticacion con Azure AD (sin login manual)', nivel=2)
lista(doc, [
    'La Canvas App detecta automaticamente el usuario logueado: User().Email, User().FullName.',
    'Al abrir la app, buscar en prode_Empleado donde email = User().Email.',
    'Si existe y tiene rol con acceso, cargar la pantalla correspondiente.',
    'Si no existe o rol = "empleado", mostrar mensaje de acceso denegado.',
    'El SuperAdmin danielgilabert@prode.es se detecta con: If(User().Email = "danielgilabert@prode.es", ...)',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '8.4 Formula de semaforo de color (replicar logica Python)', nivel=2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
run = p.add_run(
    '// En Power Apps, formula de color para cada empleado:\n'
    'Switch(\n'
    '  true,\n'
    '  ThisItem.dias_sin_fichar >= 3, RGBA(220,50,50,1),   // Rojo\n'
    '  ThisItem.dias_sin_fichar >= 1, RGBA(255,140,0,1),  // Naranja\n'
    '  ThisItem.dias_error > 0,       RGBA(0,120,215,1),  // Azul\n'
    '  RGBA(30,180,100,1)                                  // Verde\n'
    ')'
)
run.font.size = Pt(9)
run.font.name = 'Courier New'

doc.add_page_break()

# ════════════════════════════════════════
# 9. FASE 5 — POWER AUTOMATE (logica)
# ════════════════════════════════════════
seccion(doc, '9. Fase 5 — Logica de negocio con Power Automate')
parrafo(doc,
    'Power Automate reemplaza los services/*.py de Python. Cada flujo corresponde a un servicio.')
doc.add_paragraph('')

tabla(doc,
    ['Flujo Power Automate', 'Equivalente Python', 'Trigger'],
    [
        ['Calcular resumen mensual', 'calculo_service.py', 'Al pulsar boton en la app / al subir Excel'],
        ['Procesar Excel de fichajes', 'fichaje_service.py', 'Al subir fichero a SharePoint/OneDrive'],
        ['Guardar historico mensual', 'logica de guardado en resumen.py', 'Al pulsar "Guardar mes" en la app'],
        ['Generar informe PDF', 'informe_pdf_service.py', 'Al pulsar boton PDF en la app'],
        ['Generar informe Excel', 'informe_excel_service.py', 'Al pulsar boton Excel en la app'],
        ['Registrar entrada auditoria', 'registro_auditoria en services', 'En cada accion relevante'],
        ['Enviar correo notificacion', 'No implementado actualmente', 'Al fin de mes o al detectar anomalia'],
    ],
    col_widths=[2.0, 2.0, 2.0]
)
doc.add_paragraph('')

seccion(doc, '9.1 Flujo: Procesar Excel de fichajes', nivel=2)
lista(doc, [
    'Trigger: cuando se sube un archivo a la carpeta SharePoint "Fichajes/{año}/{mes}/".',
    'Accion 1: Obtener contenido del archivo Excel (conector Excel Online Business).',
    'Accion 2: Apply to each fila del Excel — normalizar nombre empleado, entrada, salida.',
    'Accion 3: Para cada fila, crear o actualizar registro en prode_FichajeRaw en Dataverse.',
    'Accion 4: Llamar al flujo "Calcular resumen mensual".',
    'Accion 5: Notificar al responsable via Teams/Outlook que el fichero fue procesado.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '9.2 Flujo: Calcular resumen mensual', nivel=2)
lista(doc, [
    'Trigger: llamado por otro flujo o por boton en la app (HTTP request).',
    'Para cada empleado del responsable: contar dias laborables del mes.',
    'Excluir festivos nacionales (lista estatica), festivos locales (prode_FestivoLocal) e incidencias.',
    'Calcular dias fichados, errores, sin fichar, horas reales vs objetivo.',
    'Guardar en prode_HistoricoResumen con fecha_proceso = hoy.',
    'Devolver JSON con el resumen calculado a la app.',
], numerada=True)
nota(doc, 'Los calculos complejos en Power Automate se hacen con expresiones. '
          'Por ejemplo: div(mul(jornada_semanal, dias_laborables), 5) para calcular horas objetivo.')

doc.add_page_break()

# ════════════════════════════════════════
# 10. FASE 6 — INFORMES
# ════════════════════════════════════════
seccion(doc, '10. Fase 6 — Informes PDF y Excel con Power Automate')
seccion(doc, '10.1 Generar PDF con plantilla Word', nivel=2)
lista(doc, [
    'Crear plantilla Word con marcadores de posicion: {{empleado_nombre}}, {{mes}}, {{dias_fichados}}, etc.',
    'Subir la plantilla a SharePoint: "Plantillas/Informe_Individual_Template.docx".',
    'Flujo Power Automate: usar el conector "Word Online Business" > "Populate a Word template".',
    'Pasar los datos del empleado como parametros.',
    'Convertir el Word relleno a PDF con el conector "OneDrive" > "Convert file" (o usar Adobe PDF).',
    'Devolver el PDF a la app via HTTP response para que el usuario lo descargue.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '10.2 Generar Excel con plantilla', nivel=2)
lista(doc, [
    'Crear plantilla Excel con formato corporativo PRODE.',
    'Flujo: usar conector "Excel Online Business" > "Add a row into a table" para rellenar datos.',
    'O usar "Create file" con contenido binario generado por script Azure Functions si se requiere formato avanzado.',
], numerada=True)
nota(doc, 'Para informes PDF con calidad alta y logo PRODE, considerar Azure Functions '
          'con Python/ReportLab como microservicio. Power Automate lo llamaria via HTTP.')

doc.add_page_break()

# ════════════════════════════════════════
# 11. FASE 7 — POWER BI
# ════════════════════════════════════════
seccion(doc, '11. Fase 7 — Integracion con Power BI')
parrafo(doc,
    'Una vez los datos esten en Dataverse, la integracion con Power BI es nativa y sin configuracion adicional.')
doc.add_paragraph('')

lista(doc, [
    'Power BI Desktop > Obtener datos > Dataverse.',
    'Iniciar sesion con cuenta @prode.es (misma que usa la app).',
    'Seleccionar las tablas: prode_Empleado, prode_HistoricoResumen, prode_Incidencia, prode_Auditoria.',
    'Crear el modelo de relaciones y las visualizaciones.',
    'Publicar en Power BI Service en el espacio de trabajo del grupo M365 de PRODE.',
    'Insertar el informe de Power BI en la pantalla de Historico de la Canvas App con el componente "Power BI tile".',
    'Configurar refresco automatico cada hora en el dataset de Power BI Service.',
], numerada=True)
doc.add_paragraph('')

nota(doc, 'Con Dataverse como backend, Power BI puede usar DirectQuery: los datos se refrescan en tiempo real '
          'sin programar actualizaciones manuales.')

doc.add_page_break()

# ════════════════════════════════════════
# 12. FASE 8 — SEGURIDAD
# ════════════════════════════════════════
seccion(doc, '12. Fase 8 — Seguridad, roles y permisos en Power Platform')
seccion(doc, '12.1 Roles de seguridad en Dataverse', nivel=2)
lista(doc, [
    'En make.powerapps.com > Dataverse > Roles de seguridad > + Nuevo rol.',
    'Crear roles: PRODE-SuperAdmin, PRODE-Administrador, PRODE-Responsable, PRODE-Coordinador.',
    'Para cada rol, definir permisos por tabla: Crear/Leer/Escribir/Eliminar a nivel de Organizacion o Usuario.',
    'PRODE-Empleado: sin acceso a ninguna tabla (no puede usar la app).',
    'PRODE-Responsable: Leer y Escribir solo filas donde responsable_id = usuario actual (user-level access).',
    'PRODE-Administrador: Leer y Escribir a nivel de Organizacion (todas las filas).',
    'PRODE-SuperAdmin: acceso total a todas las tablas y configuracion del entorno.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '12.2 Grupos de seguridad Azure AD', nivel=2)
lista(doc, [
    'Crear grupos en Azure AD: GRP-WorkTime-Admins, GRP-WorkTime-Responsables, GRP-WorkTime-Coordinadores.',
    'Asignar usuarios a los grupos segun su rol.',
    'En Power Platform Admin Center: asignar el grupo GRP-WorkTime-Admins al rol PRODE-Administrador.',
    'Esto permite gestionar accesos sin tocar la app: basta con anadir/quitar del grupo Azure AD.',
], numerada=True)
doc.add_paragraph('')

seccion(doc, '12.3 Row-level security (seguridad por fila)', nivel=2)
parrafo(doc,
    'Dataverse permite que cada responsable solo vea los registros de sus empleados, '
    'similar a como funciona actualmente con el campo responsable_id:')
lista(doc, [
    'En el rol PRODE-Responsable: configurar permisos de prode_Empleado a nivel "User" (solo filas donde el usuario es el owner o responsable).',
    'Usar columnas de propiedad (Owner) o relaciones de jerarquia para filtrar automaticamente.',
    'Esto elimina la necesidad de filtrar manualmente en el codigo como se hace hoy en empleado_repo.py.',
])

doc.add_page_break()

# ════════════════════════════════════════
# 13. FASE 9 — PRUEBAS
# ════════════════════════════════════════
seccion(doc, '13. Fase 9 — Pruebas y validacion')
tabla(doc,
    ['Tipo de prueba', 'Que verificar', 'Herramienta'],
    [
        ['Funcional', 'Login, calculo de resumen, semaforo, informes', 'Manual + Power Apps Test Studio'],
        ['Roles', 'Responsable no ve otros departamentos, Admin ve todo', 'Manual con cuentas de prueba'],
        ['Calculo', 'Resultado igual que Python actual con mismo Excel', 'Comparar CSV exportado de ambas versiones'],
        ['Rendimiento', 'Carga de pagina con 100+ empleados < 3 segundos', 'Power Apps Monitor'],
        ['Seguridad', 'Un responsable no puede acceder a datos de otro', 'Manual + Power Apps Monitor'],
        ['Integracion PBI', 'Datos aparecen en Power BI tras guardar resumen', 'Manual'],
        ['PDF/Excel', 'Informe generado con datos correctos y logo PRODE', 'Manual'],
    ],
    col_widths=[1.5, 2.5, 1.7]
)
doc.add_paragraph('')

parrafo(doc, 'Criterios de aceptacion:', negrita=True)
lista(doc, [
    'Todos los usuarios piloto pueden acceder con su cuenta @prode.es sin errores.',
    'El calculo de resumen mensual produce el mismo resultado que el sistema actual para el mismo Excel.',
    'Un responsable solo ve los empleados de su departamento.',
    'Los informes PDF incluyen logo, datos correctos y pie de pagina.',
    'Power BI muestra datos actualizados sin necesidad de cargar ningun Excel.',
    'La app es usable en navegador de escritorio y en tablet.',
])

doc.add_page_break()

# ════════════════════════════════════════
# 14. FASE 10 — DESPLIEGUE
# ════════════════════════════════════════
seccion(doc, '14. Fase 10 — Despliegue y publicacion')
lista(doc, [
    'Exportar la app desde entorno DEV: Apps > seleccionar > Exportar paquete.',
    'Importar en entorno PRD: Apps > Importar paquete > subir el .zip exportado.',
    'Reconectar los origenes de datos al entorno PRD (las conexiones se rompen al importar).',
    'Publicar la app: Archivo > Publicar > Publicar esta version.',
    'Compartir con los grupos de Azure AD: Apps > seleccionar > Compartir > buscar GRP-WorkTime-*.',
    'Anadir la app a Microsoft Teams: Teams > + Agregar una pestana > Power Apps > seleccionar WorkTimeAsistem PRODE.',
    'Comunicar la URL de acceso a todos los usuarios: https://apps.powerapps.com/play/[app-id]',
    'Monitorizar los primeros dias con Power Apps Monitor y revisar logs de auditoria.',
], numerada=True)
doc.add_paragraph('')

nota(doc, 'Mantener el sistema actual (Streamlit) activo durante el periodo de transicion (minimo 1 mes) '
          'para que los usuarios puedan comparar resultados y reportar discrepancias.')

doc.add_page_break()

# ════════════════════════════════════════
# 15. CRONOGRAMA
# ════════════════════════════════════════
seccion(doc, '15. Cronograma orientativo')
parrafo(doc, 'Estimacion en semanas de trabajo efectivo (no dias naturales):', negrita=True)
tabla(doc,
    ['Fase', 'Tarea', 'Semanas est.', 'Dependencia'],
    [
        ['Fase 1', 'Entorno M365 + licencias', '1', '-'],
        ['Fase 2', 'Modelo de datos Dataverse', '1', 'Fase 1'],
        ['Fase 3', 'Migracion de datos Supabase → Dataverse', '1', 'Fase 2'],
        ['Fase 4', 'Construccion pantallas Power Apps', '3-4', 'Fase 2'],
        ['Fase 5', 'Flujos Power Automate (logica)', '2-3', 'Fase 4'],
        ['Fase 6', 'Informes PDF/Excel con plantillas', '1-2', 'Fase 5'],
        ['Fase 7', 'Integracion Power BI', '0.5', 'Fase 2'],
        ['Fase 8', 'Seguridad y roles Dataverse', '1', 'Fases 2+4'],
        ['Fase 9', 'Pruebas y validacion', '1-2', 'Todas'],
        ['Fase 10', 'Despliegue produccion', '0.5', 'Fase 9'],
        ['TOTAL', '', '11-15 semanas', ''],
    ],
    col_widths=[0.8, 2.5, 1.2, 1.5]
)
nota(doc, 'Si se contrata un consultor Power Apps con experiencia, las fases 4 y 5 pueden reducirse a la mitad.')

doc.add_page_break()

# ════════════════════════════════════════
# 16. COSTES
# ════════════════════════════════════════
seccion(doc, '16. Costes estimados')
tabla(doc,
    ['Concepto', 'Coste aprox.', 'Frecuencia', 'Notas'],
    [
        ['Power Apps per user (si no incluido en M365)', '20 EUR/usuario/mes', 'Mensual', 'Para usuarios que usan la app. Empleados sin acceso: 0 EUR.'],
        ['Dataverse storage (si se supera el plan)', 'Incluido en Power Apps', '-', '2 GB incluidos. Adicional: 0.05 EUR/GB/mes aprox.'],
        ['Power BI Pro (para publicar dashboards)', '9 EUR/usuario/mes', 'Mensual', 'Solo para quienes publican y ven informes en Power BI Service.'],
        ['Desarrollo / consultoria', '2.000-8.000 EUR', 'Una vez', 'Segun si se hace internamente o con consultor externo.'],
        ['Formacion usuarios', '500-1.000 EUR', 'Una vez', 'Sesion formativa para responsables y administradores.'],
        ['Mantenimiento anual', '500-2.000 EUR', 'Anual', 'Actualizaciones, nuevas funcionalidades menores.'],
    ],
    col_widths=[2.0, 1.5, 1.0, 2.0]
)
nota(doc, 'Verificar con Microsoft si las licencias M365 Business Standard ya incluyen Power Apps segun el plan contratado por PRODE. '
          'Muchos planes incluyen Power Apps para uso basico (Standard connectors, SharePoint backend).')

doc.add_page_break()

# ════════════════════════════════════════
# 17. RIESGOS
# ════════════════════════════════════════
seccion(doc, '17. Riesgos y mitigaciones')
tabla(doc,
    ['Riesgo', 'Probabilidad', 'Impacto', 'Mitigacion'],
    [
        ['Licencias insuficientes (sin Power Apps Premium)', 'Media', 'Alto', 'Verificar plan M365 antes de iniciar. Alternativa: SharePoint como backend (incluido en M365).'],
        ['Perdida de datos en migracion', 'Baja', 'Alto', 'Hacer backup completo de Supabase. Migrar primero en DEV. Validar conteos de registros.'],
        ['Diferencias en calculo de resumen vs Python', 'Media', 'Medio', 'Ejecutar ambos sistemas en paralelo 1-2 meses. Comparar resultados con el mismo Excel.'],
        ['Resistencia al cambio de usuarios', 'Media', 'Medio', 'Involucrar a 2-3 responsables como usuarios piloto desde el inicio. Formacion previa al despliegue.'],
        ['Rendimiento bajo con muchos empleados', 'Baja', 'Medio', 'Power Apps + Dataverse escala bien hasta miles de registros. Revisar delegacion de consultas.'],
        ['Dependencia de consultores externos', 'Baja', 'Bajo', 'Documentar todo. Transferir conocimiento al equipo IT interno de PRODE.'],
    ],
    col_widths=[1.8, 1.0, 0.9, 2.8]
)

doc.add_page_break()

# ════════════════════════════════════════
# 18. GLOSARIO
# ════════════════════════════════════════
seccion(doc, '18. Glosario')
tabla(doc,
    ['Termino', 'Definicion'],
    [
        ['Power Apps', 'Plataforma Microsoft de desarrollo low-code para aplicaciones empresariales.'],
        ['Canvas App', 'Tipo de Power Apps donde se diseña la UI libremente como en un lienzo.'],
        ['Dataverse', 'Base de datos relacional de Microsoft integrada en Power Platform. Reemplaza a Supabase.'],
        ['Power Automate', 'Motor de flujos de trabajo y logica de negocio de Microsoft. Reemplaza services/*.py.'],
        ['Power BI', 'Herramienta de analisis e inteligencia de negocio de Microsoft.'],
        ['Azure AD', 'Servicio de identidad de Microsoft. Gestiona usuarios, grupos y autenticacion @prode.es.'],
        ['Conector', 'Adaptador en Power Apps/Automate que conecta con servicios externos (SharePoint, Excel, etc.).'],
        ['Delegacion', 'En Power Apps, capacidad de una formula para que el servidor filtre los datos (no el cliente). Importante para rendimiento con muchos registros.'],
        ['Row-level security', 'Restriccion de acceso a filas especificas de una tabla segun el usuario que consulta.'],
        ['Single Source of Truth', 'Principio de tener un unico lugar donde viven los datos, sin copias ni duplicados.'],
        ['DirectQuery', 'Modo de Power BI que lee datos en tiempo real sin importarlos localmente.'],
        ['SSO', 'Single Sign-On: autenticacion unica. El usuario se autentica una sola vez con su cuenta corporativa.'],
    ],
    col_widths=[1.8, 4.5]
)
doc.add_paragraph('')

separador(doc)
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_pie.add_run(
    f'Documento generado el {datetime.datetime.now().strftime("%d/%m/%Y")} — '
    'Desarrollado por Daniel Gilabert Cantero para Fundacion PRODE\n'
    'WorkTimeAsistem PRODE v1.0'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.save('Guia_PowerApps_WorkTimeAsistem.docx')
print('Guia Power Apps OK')
