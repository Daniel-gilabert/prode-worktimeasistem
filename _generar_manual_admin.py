from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

def titulo_seccion(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    return h

def parrafo(doc, texto, negrita=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = negrita
    return p

def tabla_roles(doc):
    tabla = doc.add_table(rows=6, cols=4)
    tabla.style = 'Table Grid'
    headers = ['Rol', 'Acceso App', 'Vista', 'Puede hacer']
    for i, h in enumerate(headers):
        cell = tabla.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    datos = [
        ['SuperAdministrador', 'Si', 'Todo + gestion app', 'Panel de control, roles, jerarquia, auditoria, todo lo demas'],
        ['Administrador', 'Si', 'Toda la entidad', 'Ver todos los departamentos, resumen global, historico global'],
        ['Responsable', 'Si', 'Su departamento', 'Ver su departamento, configurar jornadas/festivos/incidencias, exportar informes'],
        ['Coordinador', 'Si', 'Su departamento', 'Ver su departamento, exportar informes (sin configuracion avanzada)'],
        ['Empleado', 'No', 'Sin acceso', 'No puede iniciar sesion en la aplicacion'],
    ]
    for i, fila in enumerate(datos):
        for j, val in enumerate(fila):
            tabla.rows[i+1].cells[j].text = val
    doc.add_paragraph('')

# ═══════════════════════════════════════════
doc.add_heading('Manual de Usuario — WorkTimeAsistem PRODE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Version 1.0 — {datetime.datetime.now().strftime("%d/%m/%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Desarrollado por Daniel Gilabert Cantero para Fundacion PRODE').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# 1. INTRODUCCION
titulo_seccion(doc, '1. Introduccion')
parrafo(doc, 'WorkTimeAsistem PRODE es la herramienta oficial de control horario de Fundacion PRODE. '
             'Permite analizar fichajes mensuales, calcular cumplimiento horario por empleado, '
             'gestionar jornadas, festivos e incidencias, y generar informes individuales y globales.')
parrafo(doc, 'URL de acceso: https://worktime-asisten.streamlit.app/')
doc.add_paragraph('')

# 2. ROLES
titulo_seccion(doc, '2. Roles y permisos')
parrafo(doc, 'La aplicacion tiene 5 niveles de acceso. Cada usuario accede con su correo corporativo @prode.es.')
doc.add_paragraph('')
tabla_roles(doc)

# 3. COMO ACCEDER
titulo_seccion(doc, '3. Como acceder a la aplicacion')
parrafo(doc, 'Paso 1: Abre el navegador y ve a https://worktime-asisten.streamlit.app/')
parrafo(doc, 'Paso 2: Escribe tu correo corporativo (@prode.es) en el campo de login.')
parrafo(doc, 'Paso 3: Pulsa "Entrar". La aplicacion cargara automaticamente la vista correspondiente a tu rol.')
parrafo(doc, 'IMPORTANTE: Solo los correos con rol Coordinador, Responsable, Administrador o SuperAdministrador pueden acceder. '
             'Los empleados sin rol asignado no tienen acceso.')
doc.add_paragraph('')

# 4. SUPERADMINISTRADOR
titulo_seccion(doc, '4. SuperAdministrador (danielgilabert@prode.es)')
parrafo(doc, 'El SuperAdministrador tiene acceso exclusivo al Panel de Control desde el boton en el menu lateral.')
titulo_seccion(doc, '4.1 Panel de Control — Roles y acceso', nivel=2)
parrafo(doc, '- Lista todos los empleados con su rol actual y departamento asignado.')
parrafo(doc, '- Filtros por rol, departamento y nombre para encontrar rapido a cualquier persona.')
parrafo(doc, '- Para cambiar el rol de un empleado: abrir su ficha, seleccionar nuevo rol, asignar departamento y pulsar Guardar.')
parrafo(doc, '- El campo "Departamento" debe escribirse IDENTICO para todos los empleados del mismo departamento '
             '(ej: "Casa de Acogida de Cordoba"). El selector "Copiar dept" ayuda a no cometer errores.')
parrafo(doc, '- Roles disponibles: empleado / coordinador / responsable / administrador')
titulo_seccion(doc, '4.2 Panel de Control — Jerarquia', nivel=2)
parrafo(doc, '- Muestra el arbol de quien reporta a quien.')
parrafo(doc, '- Permite reasignar el responsable directo de cualquier empleado.')
parrafo(doc, '- La jerarquia es multinivel: un responsable puede tener coordinadores que a su vez tienen empleados.')
titulo_seccion(doc, '4.3 Panel de Control — Accesos al panel', nivel=2)
parrafo(doc, '- Gestiona que correos tienen acceso al panel de semaforo por departamento.')
parrafo(doc, '- Anadir o eliminar correos autorizados.')
doc.add_paragraph('')

# 5. ADMINISTRADOR
titulo_seccion(doc, '5. Administrador')
parrafo(doc, 'El administrador ve toda la entidad. Su flujo de trabajo es identico al del Responsable '
             'pero con vision global de todos los departamentos.')
parrafo(doc, '- Puede cargar el Excel mensual de fichajes.')
parrafo(doc, '- Ve el resumen de TODOS los empleados de la entidad.')
parrafo(doc, '- Accede al historico y graficas globales.')
parrafo(doc, '- Puede generar informes PDF y Excel de cualquier empleado o del global.')
doc.add_paragraph('')

# 6. RESPONSABLE Y COORDINADOR
titulo_seccion(doc, '6. Responsable y Coordinador')
titulo_seccion(doc, '6.1 Cargar el Excel de fichajes', nivel=2)
parrafo(doc, 'Paso 1: En la barra lateral, sube el archivo Excel exportado del sistema de fichajes.')
parrafo(doc, 'Paso 2: La aplicacion detecta automaticamente el mes y ano del fichero.')
parrafo(doc, 'Paso 3: Se muestra el resumen mensual de tu departamento.')
titulo_seccion(doc, '6.2 Resumen mensual', nivel=2)
parrafo(doc, '- Tabla con todos los empleados del departamento.')
parrafo(doc, '- Columnas: dias laborables, fichados, errores, sin fichar, horas reales, objetivo, diferencia.')
parrafo(doc, '- Semaforo de color por empleado (verde / azul / naranja / rojo).')
parrafo(doc, '- Botones PDF y Excel individuales al lado de cada nombre.')
titulo_seccion(doc, '6.3 Configuracion', nivel=2)
parrafo(doc, '- Jornadas: modifica las horas semanales de cada empleado.')
parrafo(doc, '- Festivos locales: anade festivos especificos para tu departamento.')
parrafo(doc, '- Incidencias: registra vacaciones, bajas y permisos por empleado y rango de fechas.')
titulo_seccion(doc, '6.4 Informes', nivel=2)
parrafo(doc, '- PDF individual: boton junto a cada nombre en el resumen.')
parrafo(doc, '- Excel individual: boton junto a cada nombre en el resumen.')
parrafo(doc, '- Los informes globales estan al final de la pagina.')
titulo_seccion(doc, '6.5 Historico y evolucion', nivel=2)
parrafo(doc, '- Pulsa "Guardar [mes]/[ano]" para guardar el resumen del mes en el historico.')
parrafo(doc, '- Las graficas muestran la evolucion del semaforo mes a mes.')
parrafo(doc, '- Tab "Horas extra": evolucion de horas extra acumuladas por mes.')
parrafo(doc, '- Archivos Excel guardados: permite almacenar y descargar los Excel mensuales.')
doc.add_paragraph('')

# 7. SEMAFORO
titulo_seccion(doc, '7. Panel de semaforo por departamento')
parrafo(doc, 'Accesible para los correos autorizados desde el boton "Panel por departamento" en el menu lateral.')
parrafo(doc, 'Muestra una tarjeta por departamento con 4 indicadores:')
parrafo(doc, '  Verde: empleados con fichaje completo y sin errores.')
parrafo(doc, '  Azul: fichaje completo pero con errores de formato.')
parrafo(doc, '  Naranja: 1-2 dias sin fichar.')
parrafo(doc, '  Rojo: 3 o mas dias sin fichar.')
parrafo(doc, 'El desplegable "Ver detalle" muestra el listado completo del departamento con estado individual.')
parrafo(doc, 'Vista "Mi departamento": muestra solo el tuyo.')
parrafo(doc, 'Vista "Todos los departamentos": vision global (requiere autorizacion).')
doc.add_paragraph('')

# 8. SEMAFORO COLORES
titulo_seccion(doc, '8. Significado de los colores en el resumen mensual')
tabla2 = doc.add_table(rows=5, cols=3)
tabla2.style = 'Table Grid'
for i, h in enumerate(['Color', 'Significado', 'Criterio']):
    tabla2.rows[0].cells[i].text = h
    tabla2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
filas2 = [
    ['Verde', 'Fichaje correcto', 'Sin dias sin fichar y sin errores'],
    ['Azul', 'Fichaje con errores', 'Sin dias sin fichar pero con errores de formato o entrada/salida'],
    ['Naranja', '1-2 dias sin fichar', 'Entre 1 y 2 dias laborables sin registro de fichaje'],
    ['Rojo', '3+ dias sin fichar', '3 o mas dias laborables sin registro de fichaje'],
]
for i, f in enumerate(filas2):
    for j, v in enumerate(f):
        tabla2.rows[i+1].cells[j].text = v
doc.add_paragraph('')

# 9. PREGUNTAS FRECUENTES
titulo_seccion(doc, '9. Preguntas frecuentes')
parrafo(doc, 'P: No puedo acceder con mi correo corporativo.', negrita=True)
parrafo(doc, 'R: Contacta con danielgilabert@prode.es para que asigne tu rol en el Panel de Control.')
doc.add_paragraph('')
parrafo(doc, 'P: Un empleado aparece en "Sin departamento asignado".', negrita=True)
parrafo(doc, 'R: El SuperAdministrador debe asignar departamento al empleado desde Panel de Control > Roles y acceso.')
doc.add_paragraph('')
parrafo(doc, 'P: El Excel no se procesa correctamente.', negrita=True)
parrafo(doc, 'R: Asegurate de exportar el fichero en el formato habitual del sistema de fichajes. '
             'Si hay nombres en el Excel que no estan en la BD, la app te ofrecera anadir los empleados nuevos.')
doc.add_paragraph('')
parrafo(doc, 'P: Como borro un Excel del historico.', negrita=True)
parrafo(doc, 'R: En la seccion "Archivos Excel guardados" dentro de Historico, pulsa el icono de papelera junto al archivo.')

doc.save('Manual_Administracion_WorkTimeAsistem.docx')
print('Manual OK')
