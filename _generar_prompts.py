from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

titulo = doc.add_heading('Prompts de continuacion - WorkTimeAsistem PRODE', 0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(f'Generado: {datetime.datetime.now().strftime("%d/%m/%Y %H:%M")}')
doc.add_paragraph('IA utilizada: Claude Sonnet 4.6 (Anthropic) ejecutado dentro de Verdent')
doc.add_paragraph('')

doc.add_heading('Prompt para ChatGPT', level=1)
doc.add_paragraph(
    'Soy Daniel Gilabert, estoy desarrollando una app llamada WorkTimeAsistem PRODE\n'
    'para Fundacion PRODE (control horario interno).\n\n'
    'Stack: Python, Streamlit, Supabase (PostgreSQL), ReportLab, OpenPyXL.\n'
    'Repositorio GitHub: https://github.com/Daniel-gilabert/prode-worktimeasistem\n'
    'Deploy: https://worktime-asisten.streamlit.app/\n\n'
    'Estructura de carpetas:\n'
    '- models/ (Empleado con campos: id, apellidos_y_nombre, email, activo, rol,\n'
    '  departamento, es_responsable, es_admin, jornada_semanal, responsable_id)\n'
    '- repositories/ (empleado_repo, festivo_repo, incidencia_repo, historico_repo,\n'
    '  departamento_repo, panel_acceso_repo, auditoria_repo)\n'
    '- services/ (auth_service, calculo_service, fichaje_service,\n'
    '  informe_pdf_service, informe_excel_service)\n'
    '- ui/ (login, resumen, configuracion, exportacion, historico,\n'
    '  panel_responsables, panel_control)\n\n'
    'Roles (campo "rol" en tabla empleados):\n'
    '- superadministrador: danielgilabert@prode.es, acceso total + gestion app\n'
    '- administrador: ve toda la entidad\n'
    '- responsable: ve solo su departamento\n'
    '- coordinador: ve solo su departamento\n'
    '- empleado: sin acceso a la app\n\n'
    'El departamento se asigna en el campo "departamento" de cada empleado.\n'
    'El semaforo agrupa por ese campo directamente.\n\n'
    'Tablas Supabase: empleados, festivos_locales, festivos_empleado, incidencias,\n'
    'historico_resumenes, panel_acceso, auditoria, departamentos\n\n'
    'Ultimo commit: fix: quitar Rol de services/__init__.py tras refactor de roles\n\n'
    'Necesito que continues ayudandome con el desarrollo.'
)

doc.add_paragraph('')
doc.add_heading('Prompt para GitHub Copilot (Chat)', level=1)
doc.add_paragraph(
    'Proyecto: WorkTimeAsistem PRODE - app Streamlit de control horario.\n'
    'Repo: https://github.com/Daniel-gilabert/prode-worktimeasistem\n\n'
    'Contexto tecnico:\n'
    '- Python + Streamlit + Supabase + ReportLab + OpenPyXL\n'
    '- Arquitectura en capas: models / repositories / services / ui\n'
    '- Modelo central: clase Empleado con campo "rol"\n'
    '  (empleado/coordinador/responsable/administrador/superadministrador)\n'
    '  y campo "departamento" (texto libre, agrupa empleados)\n'
    '- Acceso a app: roles coordinador en adelante\n'
    '- SuperAdmin hardcodeado: danielgilabert@prode.es\n'
    '- Semaforo (panel_responsables.py) agrupa por emp.departamento directo\n'
    '- Jerarquia via responsable_id (multinivel, recursivo)\n'
    '- Historico en Supabase Storage (bucket "fichajes")\n'
    '- RLS activo, clave service_role en Streamlit Secrets\n\n'
    'Archivos clave modificados recientemente:\n'
    '- models/empleado.py\n'
    '- repositories/empleado_repo.py\n'
    '- services/auth_service.py\n'
    '- ui/panel_control.py\n'
    '- ui/panel_responsables.py\n'
    '- ui/historico.py\n\n'
    'Ayudame a continuar el desarrollo desde aqui.'
)

doc.save('Prompts_Continuacion_WorkTimeAsistem.docx')
print('OK')
