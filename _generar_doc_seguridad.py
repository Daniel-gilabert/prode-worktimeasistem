from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

AZUL_PRODE   = RGBColor(0x1a, 0x3d, 0x6e)
AZUL_CLARO   = RGBColor(0x2e, 0x6d, 0xa4)
GRIS_FONDO   = RGBColor(0xf0, 0xf4, 0xf8)
VERDE        = RGBColor(0x28, 0xa7, 0x45)
ROJO         = RGBColor(0xdc, 0x35, 0x45)
NARANJA      = RGBColor(0xfd, 0x7e, 0x14)
BLANCO       = RGBColor(0xff, 0xff, 0xff)

doc = Document()

# ── Márgenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=AZUL_PRODE):
    p    = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size  = Pt(11)
        run.font.name  = "Calibri"
    return p


def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Calibri"
        r.font.color.rgb = AZUL_PRODE
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    return p


# ═════════════════════════════════════════════════════════════════════════════
# PORTADA
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(80)
r = p.add_run("INFORME DE SEGURIDAD")
r.bold = True
r.font.size  = Pt(26)
r.font.color.rgb = AZUL_PRODE
r.font.name  = "Calibri"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("WorkTimeAsistem PRODE")
r2.font.size  = Pt(18)
r2.font.color.rgb = AZUL_CLARO
r2.font.name  = "Calibri"

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Aplicación Interna de Control Horario")
r3.font.size  = Pt(13)
r3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
r3.font.name  = "Calibri"

doc.add_paragraph()
doc.add_paragraph()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(f"Fecha de emisión: {datetime.date.today().strftime('%d/%m/%Y')}")
r4.font.size  = Pt(11)
r4.font.name  = "Calibri"
r4.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

p5 = doc.add_paragraph()
p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
r5 = p5.add_run("Elaborado por: Daniel Gilabert Cantero  ·  Fundación PRODE")
r5.font.size  = Pt(11)
r5.font.name  = "Calibri"
r5.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. DESCRIPCIÓN DE LA APLICACIÓN
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Descripción de la aplicación", level=1)
add_body(doc, (
    "WorkTimeAsistem PRODE es una aplicación web interna desarrollada con Python (Streamlit) "
    "y base de datos PostgreSQL alojada en Supabase. Su función principal es el análisis de "
    "fichajes mensuales exportados en Excel, el cálculo del cumplimiento horario por empleado "
    "y la generación de informes institucionales en formato PDF y Excel."
))
add_body(doc, (
    "El acceso a la aplicación está restringido al personal autorizado de la Fundación PRODE "
    "y opera bajo un modelo de roles que determina qué información puede ver y gestionar "
    "cada usuario."
))

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 2. MODELO DE CONTROL DE ACCESO
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Control de acceso y autenticación", level=1)
add_body(doc, (
    "El sistema aplica un proceso de autenticación en tres capas secuenciales. "
    "El acceso es denegado si cualquiera de ellas no se supera:"
))

tabla_acceso = doc.add_table(rows=4, cols=3)
tabla_acceso.alignment = WD_TABLE_ALIGNMENT.CENTER
tabla_acceso.style = "Table Grid"
anchos = [Cm(1.2), Cm(5.5), Cm(8.5)]

encabezados = ["Capa", "Comprobación", "Resultado si falla"]
for i, texto in enumerate(encabezados):
    cell = tabla_acceso.rows[0].cells[i]
    cell.width = anchos[i]
    set_cell_bg(cell, "1a3d6e")
    p = cell.paragraphs[0]
    r = p.add_run(texto)
    r.bold = True
    r.font.color.rgb = BLANCO
    r.font.size = Pt(10)
    r.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

filas = [
    ("1", "Dominio corporativo @prode.es obligatorio", "Bloqueado — registrado en auditoría"),
    ("2", "El correo debe existir en la tabla empleados de Supabase", "Denegado — registrado en auditoría"),
    ("3", "El empleado debe tener rol es_responsable=true o es_admin=true", "Denegado — registrado en auditoría"),
]
colores_fila = ["f0f4f8", "ffffff", "f0f4f8"]

for idx, (capa, comp, resultado) in enumerate(filas):
    row = tabla_acceso.rows[idx + 1]
    datos = [capa, comp, resultado]
    for j, texto in enumerate(datos):
        cell = row.cells[j]
        set_cell_bg(cell, colores_fila[idx])
        p = cell.paragraphs[0]
        r = p.add_run(texto)
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        if j == 0:
            r.bold = True
            r.font.color.rgb = AZUL_PRODE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
add_body(doc, (
    "Ningún usuario con correo externo a @prode.es puede intentar el acceso, independientemente "
    "de que conozca la URL de la aplicación. Esta restricción es la primera barrera y se aplica "
    "antes de realizar ninguna consulta a la base de datos."
))

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 3. MODELO DE ROLES
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Modelo de roles y permisos", level=1)
add_body(doc, "La aplicación define dos roles con permisos diferenciados:")

tabla_roles = doc.add_table(rows=3, cols=3)
tabla_roles.alignment = WD_TABLE_ALIGNMENT.CENTER
tabla_roles.style = "Table Grid"

enc_roles = ["Rol", "Acceso", "Permisos"]
for i, texto in enumerate(enc_roles):
    cell = tabla_roles.rows[0].cells[i]
    set_cell_bg(cell, "2e6da4")
    p = cell.paragraphs[0]
    r = p.add_run(texto)
    r.bold = True
    r.font.color.rgb = BLANCO
    r.font.size = Pt(10)
    r.font.name = "Calibri"

filas_roles = [
    ("Responsable", "Sus empleados asignados", "Ver resumen mensual, gestionar festivos e incidencias de su grupo, descargar informes individuales y globales de su grupo"),
    ("Administrador", "Todos los empleados de la organización", "Todo lo anterior más: panel global por responsable, histórico completo, registro de auditoría (solo danielgilabert@prode.es)"),
]
for idx, (rol, acc, perms) in enumerate(filas_roles):
    row = tabla_roles.rows[idx + 1]
    for j, texto in enumerate([rol, acc, perms]):
        cell = row.cells[j]
        set_cell_bg(cell, "f0f4f8" if idx == 0 else "ffffff")
        p = cell.paragraphs[0]
        r = p.add_run(texto)
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        if j == 0:
            r.bold = True
            r.font.color.rgb = AZUL_PRODE

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 4. SEGURIDAD EN BASE DE DATOS
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Seguridad en base de datos (Supabase RLS)", level=1)
add_body(doc, (
    "Todas las tablas de la aplicación tienen activada la política Row Level Security (RLS) "
    "de Supabase. Esta característica garantiza que ningún acceso directo desde el exterior "
    "—aunque se disponga de la URL y la clave anónima del proyecto— pueda leer, modificar "
    "o eliminar datos."
))

add_bullet(doc, " Solo la clave service_role (clave privada del servidor) puede operar sobre las tablas.", "Política RLS:")
add_bullet(doc, " La aplicación utiliza la clave service_role almacenada como variable de entorno segura, nunca en el código fuente.", "Clave de servidor:")
add_bullet(doc, " Las siguientes tablas tienen RLS activo: empleados, festivos_locales, festivos_empleado, incidencias, panel_acceso, historico_mensual, auditoria.", "Tablas protegidas:")

doc.add_paragraph()
add_body(doc, (
    "Esto significa que incluso si un tercero obtuviera la URL pública de la base de datos "
    "y la clave anónima (que es pública por diseño en Supabase), no podría leer ningún dato. "
    "Todas las operaciones legítimas se realizan a través del servidor de la aplicación con "
    "credenciales privadas."
))

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 5. PROTECCIÓN DE CREDENCIALES
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Protección de credenciales y variables de entorno", level=1)
add_body(doc, (
    "Las credenciales de acceso a la base de datos (SUPABASE_URL y SUPABASE_KEY) no están "
    "escritas en el código fuente de la aplicación en ningún momento."
))

add_bullet(doc, " Definidas como Secrets en Streamlit Cloud, el entorno de producción.", "En producción:")
add_bullet(doc, " Almacenadas en un archivo .env local excluido del repositorio Git mediante .gitignore.", "En desarrollo local:")
add_bullet(doc, " El archivo .env nunca ha sido subido al repositorio público ni privado de GitHub.", "Repositorio:")
add_bullet(doc, " Si la aplicación arranca sin las credenciales definidas, muestra un error y detiene toda ejecución.", "Guardia de arranque:")

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 6. SISTEMA DE AUDITORÍA
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "6. Sistema de auditoría y trazabilidad", level=1)
add_body(doc, (
    "La aplicación registra automáticamente en la tabla auditoria de Supabase todos los "
    "eventos relevantes de seguridad y operación:"
))

tabla_aud = doc.add_table(rows=5, cols=3)
tabla_aud.alignment = WD_TABLE_ALIGNMENT.CENTER
tabla_aud.style = "Table Grid"

enc_aud = ["Evento", "Detalle registrado", "Resultado posible"]
for i, texto in enumerate(enc_aud):
    cell = tabla_aud.rows[0].cells[i]
    set_cell_bg(cell, "1a3d6e")
    p = cell.paragraphs[0]
    r = p.add_run(texto)
    r.bold = True
    r.font.color.rgb = BLANCO
    r.font.size = Pt(10)
    r.font.name = "Calibri"

filas_aud = [
    ("Intento de login", "Correo, timestamp, IP", "ok / denegado / bloqueado"),
    ("Dominio no permitido", "Correo intentado, timestamp", "bloqueado"),
    ("Sin registro en BD", "Correo, motivo exacto", "denegado"),
    ("Carga de Excel", "Correo, mes/año procesado", "ok / error"),
]
for idx, (ev, det, res) in enumerate(filas_aud):
    row = tabla_aud.rows[idx + 1]
    for j, texto in enumerate([ev, det, res]):
        cell = row.cells[j]
        set_cell_bg(cell, "f0f4f8" if idx % 2 == 0 else "ffffff")
        p = cell.paragraphs[0]
        r = p.add_run(texto)
        r.font.size = Pt(10)
        r.font.name = "Calibri"

doc.add_paragraph()
add_body(doc, (
    "El acceso a la vista de auditoría dentro de la aplicación está reservado exclusivamente "
    "al usuario danielgilabert@prode.es. El resto de usuarios, incluso con rol de administrador, "
    "no pueden visualizar ni consultar el registro de auditoría."
))

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 7. SEGURIDAD EN EL DESPLIEGUE
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "7. Seguridad en el despliegue (Streamlit Cloud)", level=1)
add_body(doc, (
    "La aplicación está desplegada en Streamlit Cloud, plataforma gestionada con los siguientes "
    "controles de seguridad adicionales:"
))

add_bullet(doc, " Toda la comunicación entre el navegador del usuario y la aplicación está cifrada con TLS 1.2/1.3 (HTTPS).", "Comunicaciones cifradas (HTTPS):")
add_bullet(doc, " Solo el propietario del repositorio (danielgilabert) puede desplegar cambios en producción.", "Control de despliegue:")
add_bullet(doc, " Los Secrets de Streamlit Cloud son accesibles únicamente desde la ejecución del servidor, nunca expuestos al navegador.", "Secrets aislados:")
add_bullet(doc, " La aplicación no almacena contraseñas. La autenticación se basa en correo electrónico verificado contra la base de datos corporativa.", "Sin contraseñas en texto plano:")

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 8. RESUMEN EJECUTIVO
# ═════════════════════════════════════════════════════════════════════════════
add_heading(doc, "8. Resumen ejecutivo de controles de seguridad", level=1)

tabla_resumen = doc.add_table(rows=9, cols=3)
tabla_resumen.alignment = WD_TABLE_ALIGNMENT.CENTER
tabla_resumen.style = "Table Grid"

enc_res = ["Control", "Implementación", "Estado"]
for i, texto in enumerate(enc_res):
    cell = tabla_resumen.rows[0].cells[i]
    set_cell_bg(cell, "1a3d6e")
    p = cell.paragraphs[0]
    r = p.add_run(texto)
    r.bold = True
    r.font.color.rgb = BLANCO
    r.font.size = Pt(10)
    r.font.name = "Calibri"

controles = [
    ("Restricción de dominio",        "Solo @prode.es puede intentar acceso",              "ACTIVO"),
    ("Autenticación por BD",          "Correo verificado en tabla empleados",               "ACTIVO"),
    ("Control de roles",              "es_responsable / es_admin por empleado",             "ACTIVO"),
    ("RLS en todas las tablas",       "Bloqueo total con clave anónima",                    "ACTIVO"),
    ("Credenciales en variables env", "Nunca en código fuente ni repositorio",              "ACTIVO"),
    ("Comunicación HTTPS",            "TLS 1.2/1.3 en Streamlit Cloud",                    "ACTIVO"),
    ("Registro de auditoría",         "Tabla auditoria en Supabase, acceso restringido",   "ACTIVO"),
    ("Logging de aplicación",         "Archivo logs/app.log con rotación automática",      "ACTIVO"),
]
for idx, (ctrl, impl, estado) in enumerate(controles):
    row = tabla_resumen.rows[idx + 1]
    for j, texto in enumerate([ctrl, impl, estado]):
        cell = row.cells[j]
        set_cell_bg(cell, "f0f4f8" if idx % 2 == 0 else "ffffff")
        p = cell.paragraphs[0]
        r = p.add_run(texto)
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        if j == 0:
            r.bold = True
            r.font.color.rgb = AZUL_PRODE
        if j == 2:
            r.bold = True
            r.font.color.rgb = VERDE

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# PIE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p_pie = doc.add_paragraph()
p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_pie.paragraph_format.space_before = Pt(200)
r_pie = p_pie.add_run(
    f"Documento generado el {datetime.date.today().strftime('%d/%m/%Y')}  ·  "
    "Elaborado por Daniel Gilabert Cantero  ·  Fundación PRODE"
)
r_pie.font.size  = Pt(10)
r_pie.font.name  = "Calibri"
r_pie.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

ruta = "Seguridad_WorkTimeAsistem_PRODE.docx"
doc.save(ruta)
print(f"Documento generado: {ruta}")
