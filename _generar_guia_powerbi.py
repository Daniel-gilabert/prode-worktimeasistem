from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()

def h(doc, texto, nivel=1):
    return doc.add_heading(texto, level=nivel)

def p(doc, texto, negrita=False):
    par = doc.add_paragraph()
    run = par.add_run(texto)
    run.bold = negrita
    return par

doc.add_heading('Guia de integracion Power BI — WorkTimeAsistem PRODE', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'Version 1.0 — {datetime.datetime.now().strftime("%d/%m/%Y")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Objetivo: eliminar la carga manual del Excel y que Power BI lea directamente de Supabase').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── SECCION 1 ──
h(doc, '1. Situacion actual vs situacion objetivo')
p(doc, 'ACTUAL:', negrita=True)
p(doc, '1. Se exporta un Excel del sistema de fichajes.')
p(doc, '2. Se sube manualmente a la app.')
p(doc, '3. La app calcula y muestra el resumen.')
p(doc, '4. Para Power BI hay que exportar otro Excel.')
doc.add_paragraph('')
p(doc, 'OBJETIVO:', negrita=True)
p(doc, '1. El sistema de fichajes escribe directamente en Supabase (o la app lo hace).')
p(doc, '2. Power BI lee de Supabase en tiempo casi real.')
p(doc, '3. No hay carga manual de ningun Excel.')
doc.add_paragraph('')

# ── SECCION 2 ──
h(doc, '2. Arquitectura recomendada')
p(doc, 'Supabase (PostgreSQL) actua como unica fuente de verdad (Single Source of Truth):')
doc.add_paragraph('')
t = doc.add_table(rows=4, cols=2)
t.style = 'Table Grid'
for i, h2 in enumerate(['Componente', 'Funcion']):
    t.rows[0].cells[i].text = h2
    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
filas = [
    ['Sistema de fichajes', 'Exporta CSV/Excel o escribe via API en tabla "fichajes_raw" de Supabase'],
    ['WorkTimeAsistem PRODE', 'Lee de Supabase, calcula, guarda resumenes en "historico_resumenes"'],
    ['Power BI', 'Conecta directamente a Supabase via conector PostgreSQL — sin Excel'],
]
for i, f in enumerate(filas):
    for j, v in enumerate(f):
        t.rows[i+1].cells[j].text = v
doc.add_paragraph('')

# ── SECCION 3 ──
h(doc, '3. Opcion A (RECOMENDADA): Power BI conecta a Supabase via PostgreSQL')
p(doc, 'Esta es la opcion mas simple. Power BI Desktop tiene conector nativo PostgreSQL.', negrita=True)
doc.add_paragraph('')
p(doc, 'PASO 1 — Obtener credenciales de conexion de Supabase:', negrita=True)
p(doc, '  1. Ve a https://supabase.com/dashboard/project/gqfiarxccbaznjxispsv')
p(doc, '  2. Settings > Database > Connection string > URI')
p(doc, '  3. Copia el host, puerto, nombre de BD, usuario y contrasena.')
p(doc, '  4. Host: db.gqfiarxccbaznjxispsv.supabase.co')
p(doc, '  5. Puerto: 5432')
p(doc, '  6. Base de datos: postgres')
p(doc, '  7. Usuario: postgres')
p(doc, '  8. Contrasena: la que tienes en Settings > Database > Reset password')
doc.add_paragraph('')
p(doc, 'PASO 2 — Conectar Power BI Desktop a Supabase:', negrita=True)
p(doc, '  1. Abre Power BI Desktop.')
p(doc, '  2. Inicio > Obtener datos > Base de datos PostgreSQL.')
p(doc, '  3. Servidor: db.gqfiarxccbaznjxispsv.supabase.co')
p(doc, '  4. Base de datos: postgres')
p(doc, '  5. Modo de conectividad: DirectQuery (tiempo real) o Importar (snapshot).')
p(doc, '  6. Introduce usuario y contrasena de la BD.')
p(doc, '  7. Selecciona las tablas: empleados, historico_resumenes, incidencias, etc.')
doc.add_paragraph('')
p(doc, 'PASO 3 — Instalar driver si es necesario:', negrita=True)
p(doc, '  Si Power BI pide driver PostgreSQL, descarga e instala:')
p(doc, '  https://www.postgresql.org/ftp/odbc/versions/msi/')
p(doc, '  Busca: psqlodbc_16_x64.zip — instala y reinicia Power BI.')
doc.add_paragraph('')
p(doc, 'PASO 4 — Publicar en Power BI Service con refresco automatico:', negrita=True)
p(doc, '  1. Publica el informe en Power BI Service.')
p(doc, '  2. En el dataset, configura "Actualizar ahora" o "Actualizar programado".')
p(doc, '  3. Para DirectQuery no necesitas programar — lee en tiempo real cada vez que abres el informe.')
p(doc, '  NOTA: DirectQuery requiere Power BI Gateway si el servidor no es publico.')
p(doc, '  Supabase SI es publico por lo que NO necesitas Gateway.')
doc.add_paragraph('')

# ── SECCION 4 ──
h(doc, '4. Opcion B: API REST de Supabase (para Power BI sin driver PostgreSQL)')
p(doc, 'Si no puedes instalar el driver PostgreSQL, Power BI puede leer de Supabase via su API REST.', negrita=True)
doc.add_paragraph('')
p(doc, 'PASO 1 — En Power BI Desktop:', negrita=True)
p(doc, '  1. Inicio > Obtener datos > Web.')
p(doc, '  2. URL: https://gqfiarxccbaznjxispsv.supabase.co/rest/v1/historico_resumenes?select=*')
p(doc, '  3. En Opciones avanzadas, anadir cabeceras HTTP:')
p(doc, '     apikey: [tu SUPABASE_KEY anon]')
p(doc, '     Authorization: Bearer [tu SUPABASE_KEY anon]')
p(doc, '  4. Power BI parseara el JSON automaticamente.')
doc.add_paragraph('')
p(doc, 'Tablas disponibles via API:', negrita=True)
p(doc, '  /rest/v1/empleados?select=*&activo=eq.true')
p(doc, '  /rest/v1/historico_resumenes?select=*')
p(doc, '  /rest/v1/incidencias?select=*')
p(doc, '  /rest/v1/festivos_locales?select=*')
doc.add_paragraph('')

# ── SECCION 5 ──
h(doc, '5. Tablas clave de Supabase para Power BI')
t2 = doc.add_table(rows=6, cols=3)
t2.style = 'Table Grid'
for i, h3 in enumerate(['Tabla', 'Contenido', 'Uso en Power BI']):
    t2.rows[0].cells[i].text = h3
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
filas2 = [
    ['empleados', 'Todos los empleados con rol y departamento', 'Dimensiones: nombre, dept, rol, jornada'],
    ['historico_resumenes', 'Resumenes mensuales guardados por responsable', 'Hechos: dias fichados, horas, diferencia, errores'],
    ['incidencias', 'Vacaciones, bajas, permisos por empleado', 'Dimension de tiempo ausente'],
    ['festivos_locales', 'Festivos por responsable y ano', 'Calendario laboral'],
    ['auditoria', 'Log de accesos y acciones', 'Informe de seguridad y trazabilidad'],
]
for i, f in enumerate(filas2):
    for j, v in enumerate(f):
        t2.rows[i+1].cells[j].text = v
doc.add_paragraph('')

# ── SECCION 6 ──
h(doc, '6. Automatizar la carga de fichajes (eliminar el Excel completamente)')
p(doc, 'Para que no haga falta subir ningun Excel, hay dos caminos:', negrita=True)
doc.add_paragraph('')
p(doc, 'CAMINO 1 — El sistema de fichajes exporta a Supabase directamente:', negrita=True)
p(doc, '  Si el sistema de fichajes tiene API o puede exportar CSV automaticamente:')
p(doc, '  1. Crear tabla "fichajes_raw" en Supabase con columnas: empleado_nombre, fecha, entrada, salida.')
p(doc, '  2. Configurar el sistema de fichajes para que escriba en esa tabla via API REST de Supabase.')
p(doc, '  3. La app y Power BI leen de "fichajes_raw" sin necesidad de Excel.')
doc.add_paragraph('')
p(doc, 'CAMINO 2 — Script automatico de carga (si el sistema genera ficheros en una carpeta):', negrita=True)
p(doc, '  1. El sistema de fichajes deposita el Excel en una carpeta de red o SharePoint.')
p(doc, '  2. Un script Python (ejecutado como tarea programada) lee el Excel y lo sube a Supabase.')
p(doc, '  3. Configurar en Windows Task Scheduler para que corra cada hora o cada noche.')
p(doc, '  Script minimo (Python):')
p(doc, '''
  import pandas as pd
  from supabase import create_client
  import os

  sb = create_client(os.environ["SUPABASE_URL"], os.environ["SUPABASE_KEY"])
  df = pd.read_excel("ruta/al/fichaje_mensual.xlsx")
  # normalizar columnas...
  sb.table("fichajes_raw").upsert(df.to_dict("records")).execute()
  print("Cargado OK")
''')
doc.add_paragraph('')

# ── SECCION 7 ──
h(doc, '7. Modelo de datos sugerido para Power BI')
p(doc, 'Relaciones recomendadas en Power BI:', negrita=True)
p(doc, '  empleados[id] --- historico_resumenes[empleado_id]  (1 a muchos)')
p(doc, '  empleados[id] --- incidencias[empleado_id]          (1 a muchos)')
p(doc, '  empleados[departamento] = dimension de departamento')
doc.add_paragraph('')
p(doc, 'Medidas DAX utiles:', negrita=True)
p(doc, '  % Cumplimiento = DIVIDE(COUNTROWS(FILTER(historico_resumenes, [sin_fichar]=0)), COUNTROWS(historico_resumenes))')
p(doc, '  Horas extra total = SUM(historico_resumenes[horas_extra])')
p(doc, '  Empleados en rojo = COUNTROWS(FILTER(historico_resumenes, [sin_fichar]>=3))')
doc.add_paragraph('')

# ── SECCION 8 ──
h(doc, '8. Resumen de pasos para activar la conexion hoy mismo')
pasos = [
    '1. Abre Supabase > Settings > Database > anota host, puerto y contrasena.',
    '2. Descarga e instala el driver ODBC PostgreSQL si no lo tienes.',
    '3. En Power BI Desktop: Obtener datos > PostgreSQL.',
    '4. Conecta con DirectQuery para tiempo casi real.',
    '5. Selecciona las tablas: empleados, historico_resumenes, incidencias.',
    '6. Crea el modelo de relaciones y tus visualizaciones.',
    '7. Publica en Power BI Service.',
    '8. Comparte el enlace del informe con tu equipo.',
    '(Opcional) Configura el script de carga automatica para eliminar el Excel.',
]
for paso in pasos:
    doc.add_paragraph(paso, style='List Number')

doc.add_paragraph('')
p(doc, 'NOTA IMPORTANTE sobre seguridad:', negrita=True)
p(doc, 'Usa siempre la contrasena de base de datos (no la SUPABASE_KEY del codigo). '
       'La contrasena de BD solo da acceso a PostgreSQL y no expone la API REST. '
       'Puedes crear un usuario de solo lectura en Supabase para Power BI si deseas mayor seguridad.')

doc.save('Guia_PowerBI_WorkTimeAsistem.docx')
print('Guia PowerBI OK')
